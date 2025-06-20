from flask import Flask, request, jsonify, render_template, session, redirect, url_for, make_response
from flask_sqlalchemy import SQLAlchemy
from datetime import datetime, timedelta
from flask_cors import CORS
from dotenv import load_dotenv
from urllib.parse import urlparse
import pymysql
import math
import base64
import os
import re
import csv
import json
from dateutil.parser import parse as parse_date
import PyPDF2
import openai
from openai import OpenAI
from docx import Document
from io import BytesIO, StringIO
from weasyprint import HTML
from weasyprint.text.fonts import FontConfiguration


# Load environment variables from .env file
load_dotenv()

# Initialize Flask app
app = Flask(__name__)
CORS(app)  # Enable Cross-Origin Resource Sharing
app.secret_key = os.environ.get('SECRET_KEY')  # Set secret key for session management
client = OpenAI(api_key=os.environ.get('OPENAI_API_KEY'))  # OpenAI API


# Configure Flask app settings
app.config['ENV'] = os.environ.get('FLASK_ENV', 'production')
uri = os.environ.get('DATABASE_URL')
if uri:
    if uri.startswith("postgres://"):
        uri = uri.replace("postgres://", "mysql+pymysql://", 1)
    
    parsed = urlparse(uri)
    db_config = {
        'username': parsed.username,
        'password': parsed.password,
        'hostname': parsed.hostname,
        'port': parsed.port,
        'database': parsed.path[1:],
    }
    
    app.config['SQLALCHEMY_DATABASE_URI'] = (
        f"mysql+pymysql://{db_config['username']}:{db_config['password']}"
        f"@{db_config['hostname']}:{db_config['port']}"
        f"/{db_config['database']}"
    )

# Use this for Railway SSL
app.config['SQLALCHEMY_ENGINE_OPTIONS'] = {
    'pool_pre_ping': True,
    'pool_recycle': 300,
    'pool_size': 10,
    'max_overflow': 20,
    'connect_args': {
        'ssl': {
            'verify_cert': False,
            'verify_identity': False
        }
    }
}


app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db = SQLAlchemy(app)

# Admin credentials for dashboard access
ADMIN_EMAIL = os.environ.get('ADMIN_EMAIL', 'admin@aibidmaster.com')
ADMIN_PASSWORD = os.environ.get('ADMIN_PASSWORD', 'Admin@1235')

# Virginia-specific constants (2025, VDOT and RSMeans standards)
ASPHALT_DENSITY = 145  # lbs per cubic foot, VDOT average for hot-mix asphalt (HMA)
RECYCLED_DENSITY = 140  # lbs per cubic foot for recycled asphalt
CONCRETE_DENSITY = 150  # lbs per cubic foot
LABOR_RATE = 62.50  # $/hour, Virginia union labor rates, 2025
MATERIAL_MARKUP = 1.15  # 15% markup, standard for Virginia contractors
EQUIPMENT_RATE_MULTIPLIER = 1.12  # 12% markup for equipment overhead
PROFIT_MARGIN = 0.10  # 10% profit margin, competitive for VDOT projects
OVERHEAD_RATE = 0.12  # 12% overhead, typical for Virginia construction


# Material unit costs (Virginia market rates, 2025)
MATERIAL_UNIT_COSTS = {
    'asphalt': 110,                # $/ton
    'recycled asphalt': 85,        # $/ton
    'concrete': 170,               # $/cubic yard
    'bituminous surface': 120,     # $/ton
    'sealcoat': 0.55,              # $/sq ft
    'rebar': 0.80,                 # $/lb
    'aggregate base': 42           # $/ton
}

# Material thickness standards (in feet)
THICKNESS = {
    'asphalt': 0.33,          # 4 inches
    'recycled asphalt': 0.25, # 3 inches
    'bituminous surface': 0.17, # 2 inches
    'concrete': 0.42,         # 6 inches
    'sealcoat': 0.02          # 0.25 inches
}

MATERIAL_CONSTANTS = {
    'asphalt': {'density': 145, 'thickness': 0.33},  # VDOT standard HMA
    'recycled asphalt': {'density': 140, 'thickness': 0.25},  # Virginia recycled mix
    'concrete': {'density': 150, 'thickness': 0.42},  # VDOT concrete spec
    'bituminous surface': {'density': 145, 'thickness': 0.17},  # Virginia surface treatment
    'sealcoat': {'density': 100, 'thickness': 0.02}   # Sealcoat
}


# Project Model: Defines the database schema for storing project details
class Project(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(500), nullable=False)  # Project name
    type = db.Column(db.String(100), nullable=False)  # Project type (e.g., road)
    location = db.Column(db.String(255), nullable=False)  # Project location
    submitted = db.Column(db.Date, nullable=False)  # Submission date
    status = db.Column(db.Enum('pending', 'accepted', 'rejected'), default='pending')  # Project status
    cost = db.Column(db.String(50), nullable=False)  # Estimated cost
    completion_date = db.Column(db.Date)  # Completion date
    land_mile = db.Column(db.Float)  # Length in lane miles
    width = db.Column(db.Float)  # Width in feet
    area = db.Column(db.Float, nullable=False)  # Area in sq ft
    material = db.Column(db.String(50), nullable=False)  # Material type (e.g., concrete)
    tonnage = db.Column(db.Float)  # Material tonnage
    scope = db.Column(db.Text, nullable=False)  # Project scope
    requirements = db.Column(db.Text)  # Special requirements
    estimated_cost = db.Column(db.String(50))  # Estimated cost (formatted)
    profit_margin = db.Column(db.String(20))  # Profit margin percentage
    success_probability = db.Column(db.String(20))  # Success probability
    asphalt_tons = db.Column(db.Float)  # Asphalt quantity in tons
    concrete_yds = db.Column(db.Float)  # Concrete quantity in cubic yards
    bituminous_tons = db.Column(db.Float)
    sealcoat_sqft = db.Column(db.Float)
    rebar_lbs = db.Column(db.Float)  # Rebar quantity in pounds
    aggregate_tons = db.Column(db.Float)  # Aggregate quantity in tons
    management_hours = db.Column(db.Integer)  # Management labor hours
    prep_hours = db.Column(db.Integer)  # Site preparation labor hours
    paving_hours = db.Column(db.Integer)  # Paving labor hours
    finishing_hours = db.Column(db.Integer)  # Finishing labor hours



# Serve the main index page
@app.route('/')
def index():
    return render_template('index.html')


# Add headers to prevent caching for admin sessions
@app.after_request
def add_header(response):
    if 'admin_logged_in' in session:
        response.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0'
        response.headers['Pragma'] = 'no-cache'
        response.headers['Expires'] = '0'
    return response


# Serve the admin login page
@app.route('/admin', methods=['GET'])
def admin_login_page():
    session.pop('admin_logged_in', None)  # Clear admin session
    return render_template('admin_login.html')


# Handle admin login
@app.route('/admin/login', methods=['POST'])
def admin_login():
    data = request.form
    email = data.get('email')
    password = data.get('password')
    
    # Verify admin credentials
    if email == ADMIN_EMAIL and password == ADMIN_PASSWORD:
        session['admin_logged_in'] = True
        return jsonify({'success': True})
    
    return jsonify({'success': False, 'message': 'Invalid credentials, Please Try Again'})


# Serve the admin dashboard
@app.route('/admin/dashboard', methods=['GET'])
def admin_dashboard():
    if not session.get('admin_logged_in'):
        return redirect(url_for('admin_login_page'))
    
    status = request.args.get('status', 'pending')  # Filter by project status
    query = Project.query
    
    if status != 'all':
        query = query.filter_by(status=status)
    
    projects = query.all()
    return render_template('admin_dashboard.html', 
                           status=status,
                           projects=projects)


# Serve project detail page
@app.route('/admin/projects/<int:project_id>', methods=['GET'])
def admin_project_detail(project_id):
    if not session.get('admin_logged_in'):
        return redirect(url_for('admin_login_page'))
    
    project = db.session.get(Project, project_id)
    if not project:
        return redirect(url_for('admin_dashboard'))
    
    return render_template('project_detail.html', project=project)


# Handle admin logout
@app.route('/admin/logout', methods=['GET'])
def admin_logout():
    session.pop('admin_logged_in', None)
    return redirect(url_for('admin_login_page'))


# Get all projects details
@app.route('/api/admin/projects', methods=['GET'])
def get_projects():
    if not session.get('admin_logged_in'):
        return jsonify({'error': 'Unauthorized'}), 401
    
    status = request.args.get('status', 'pending')
    
    if status == 'all':
        projects = Project.query.all()
    else:
        projects = Project.query.filter_by(status=status).all()
    
    # Return project data as JSON
    return jsonify([{
        'id': p.id,
        'name': p.name,
        'type': p.type,
        'location': p.location,
        'submitted': p.submitted.strftime('%Y-%m-%d'),
        'status': p.status,
        'cost': p.cost,
    } for p in projects])


# Accept project
@app.route('/api/admin/projects/<int:project_id>/accept', methods=['POST'])
def accept_project(project_id):
    project = db.session.get(Project, project_id)
    if project:
        project.status = 'accepted'
        db.session.commit()
        return jsonify({'message': 'Project accepted'})
    return jsonify({'error': 'Project not found'}), 404


# Reject project
@app.route('/api/admin/projects/<int:project_id>/reject', methods=['POST'])
def reject_project(project_id):
    project = db.session.get(Project, project_id)
    if project:
        project.status = 'rejected'
        db.session.commit()
        return jsonify({'message': 'Project rejected'})
    return jsonify({'error': 'Project not found'}), 404


# Delete project
@app.route('/api/admin/projects/<int:project_id>', methods=['DELETE'])
def delete_project(project_id):
    project = db.session.get(Project, project_id)
    if project:
        db.session.delete(project)
        db.session.commit()
        return jsonify({'message': 'Project deleted'})
    return jsonify({'error': 'Project not found'}), 404


# Get project details
@app.route('/api/admin/projects/<int:project_id>', methods=['GET'])
def get_project(project_id):
    project = db.session.get(Project, project_id)
    if project:
        # Return detailed project data as JSON
        return jsonify({
            'id': project.id,
            'name': project.name,
            'type': project.type,
            'location': project.location,
            'submitted': project.submitted.strftime('%Y-%m-%d'),
            'status': project.status,
            'cost': project.cost,
            'details': {
                'completionDate': project.completion_date.strftime('%Y-%m-%d') if project.completion_date else None,
                'landMile': project.land_mile,
                'width': project.width,
                'area': project.area,
                'material': project.material,
                'tonnage': project.tonnage,
                'scope': project.scope,
                'requirements': project.requirements,
                'estimatedCost': project.estimated_cost,
                'profitMargin': project.profit_margin,
                'successProbability': project.success_probability,
                'asphalt': project.asphalt_tons,
                'concrete': project.concrete_yds,
                'rebar': project.rebar_lbs,
                'aggregate': project.aggregate_tons,
                'managementHours': project.management_hours,
                'prepHours': project.prep_hours,
                'pavingHours': project.paving_hours,
                'finishingHours': project.finishing_hours
            }
        })
    return jsonify({'error': 'Project not found'}), 404


# Extract text from PDF files
def extract_text_from_pdf(file):
    pdf_reader = PyPDF2.PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text

# Extract text from DOCX files
def extract_text_from_docx(file):
    doc = Document(BytesIO(file.read()))
    text = ""
    for para in doc.paragraphs:
        if para.text.strip():
            text += para.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    text += cell_text + "\n"
    if not text.strip():
        raise ValueError("No text found in DOCX.")
    return text

# Extract RFP data using regex patterns
def extract_rfp_data(text):
    """Extract project details from RFP text using regex patterns."""
    data = {}
    
    # Normalize text for consistent matching
    normalized_text = re.sub(r'\s+', ' ', text).lower()
    original_text = text  # Preserve original for section extraction
    
    # Define regex patterns for key fields
    patterns = [
        # Project Name
        (r'(?:project\s*(?:name|title|description)|job\s*(?:name|title))[:\s]*([^\n;]+)', 'project_name'),
        (r'rfp\s*[#№]\s*[\w-]+\s*[-–—:]\s*([^\n;]+)', 'project_name'),
        # Project Location
        (r'(?:project\s*location|location|place|site)[:\s]*([^\n;]+)', 'project_location'),
        (r'in\s*([^\n,]+)(?:\s*(?:county|city|state|subdivision))', 'project_location'),
        # Completion Date
        (r'(?:completion\s*date|target\s*date|work\s*(?:must\s*be\s*)?completed\s*by|deadline)[:\s]*([a-z]+\s*\d{1,2},\s*\d{4}|\d{4}-\d{2}-\d{2})', 'completion_date'),
        (r'fully\s*completed\s*by\s*([a-z]+\s*\d{1,2},\s*\d{4}|\d{4}-\d{2}-\d{2})', 'completion_date'),
        # Project Duration
        (r'(?:duration|project\s*duration|timeline)\s*(?:\(?\s*weeks?\s*\))?[:\s]*(\d+)', 'project_duration'),
        # Lane Mile
        (r'(\d+(?:\.\d+)?)\s*(?:lane\s*[-–—]?\s*mi(?:les?)?|mi(?:les?)?)', 'land_mile'),
        # Width
        (r'(\d+(?:\.\d+)?)\s*(?:ft|feet|foot)(?:\s*width)?', 'width'),
        # Area
        (r'(?:area\s*\(?\s*sq\s*ft\s*\)?|square\s*footage)[:\s]*([\d,]+(?:\.\d+)?)', 'project_area'),
        (r'(\d+,?\d*)\s*(?:ft²|square\s*feet|sq\s*ft)', 'project_area'),
        # Material Type (capture multiple occurrences)
        (r'\b(asphalt|hma|wma|concrete|aggregate\s*base|recycled\s*asphalt|bituminous\s*surface|subbase|geotextile|sealcoat|thermoplastic\s*striping|curb|sidewalk|pavers|drainage\s*pipe|stormwater\s*structure)\b', 'material_type'),
        # Tonnage
        (r'(?:tonnage|quantity\s*tons?)[:\s]*([\d,]+(?:\.\d+)?)\s*(?:tons?)', 'tonnage'),
        # Quantities with units
        (r'(\d+,?\d*(?:\.\d+)?)\s*(ft²|ft³|yd³|tons?|lbs?|ft|square\s*feet|cubic\s*yards|linear\s*feet|sq\s*ft|each)\s*(?:of\s*)?(asphalt|hma|concrete|aggregate\s*base|rebar|curb|sidewalk|pavers|drainage\s*pipe|stormwater\s*structure)', 'quantities'),
    ]
    
    # Extract fields using patterns
    quantities = []
    for pattern, key in patterns:
        if key == 'quantities':
            matches = re.finditer(pattern, normalized_text, re.IGNORECASE)
            for match in matches:
                qty = match.group(1).replace(',', '')
                unit = match.group(2).lower()
                material = match.group(3).lower()
                quantities.append({'quantity': qty, 'unit': unit, 'material': material})
        else:
            if key not in data:
                match = re.search(pattern, normalized_text, re.IGNORECASE)
                if match:
                    data[key] = match.group(1).strip()
                    if key in ['land_mile', 'width', 'tonnage', 'project_area']:
                        data[key] = data[key].replace(',', '')
    
    # Process quantities
    if quantities:
        data['quantities'] = []
        for q in quantities:
            try:
                qty = float(q['quantity'])
                unit = q['unit'].replace('square feet', 'sq ft').replace('cubic yards', 'yd³').replace('linear feet', 'ft')
                material = q['material']
                if material == 'hma':
                    material = 'asphalt'
                data['quantities'].append({'quantity': qty, 'unit': unit, 'material': material})
                # Assign primary material and tonnage if applicable
                if material in ['asphalt', 'aggregate base'] and unit in ['tons']:
                    data['tonnage'] = str(qty)
                    data['material_type'] = material
                elif material == 'concrete' and unit in ['yd³']:
                    data['concrete_yds'] = qty
                    data['material_type'] = material
            except ValueError:
                pass
    
    # Calculate area if not provided
    if 'project_area' not in data and 'land_mile' in data and 'width' in data:
        try:
            land_mile = float(data['land_mile'])
            width = float(data['width'])
            data['project_area'] = str(round((land_mile * 5280) * width))
        except (ValueError, TypeError):
            pass
    
    # Extract sections like scope and requirements
    section_patterns = [
        ('project_scope', r'(?:scope\s*of\s*work|project\s*description|work\s*details)[:\s]*'),
        ('project_requirements', r'(?:special\s*(?:conditions|notes|requirements)|additional\s*notes)[:\s]*')
    ]
    
    for key, pattern in section_patterns:
        if key not in data:
            match = re.search(pattern, normalized_text, re.IGNORECASE)
            if match:
                start_pos = match.end()
                # Find section end (next section or end of text)
                end_pos = len(normalized_text)
                for end_pattern in [r'\n\s*\n', r'\n[A-Z][A-Z\s]+[:\s]']:
                    end_match = re.search(end_pattern, normalized_text[start_pos:], re.IGNORECASE)
                    if end_match:
                        end_pos = min(end_pos, start_pos + end_match.start())
                # Use original text to preserve formatting
                section_text = original_text[original_text.lower().find(normalized_text[start_pos:end_pos]):].strip()
                if section_text:
                    data[key] = section_text[:1000]  # Limit length to avoid DB issues
    
    # Clean extracted data
    for key in data:
        if isinstance(data[key], str):
            data[key] = re.sub(r'^[:;,.]+|[:;,.]+$', '', data[key].strip())
            if key in ['project_name', 'project_location', 'material_type']:
                data[key] = data[key][0].upper() + data[key][1:] if data[key] else ''
    
    return data

# Extract RFP data using OpenAI GPT
def extract_fields_with_openai(text):
    """Use OpenAI GPT to extract structured data from RFP text."""
    
    # Define prompt with examples and explicit instructions
    prompt = """
You are an expert at extracting structured data from construction RFPs. Extract and map all relevant fields from the provided RFP text to the following schema, even if the field names in the RFP differ or are in a different format. Use synonymous terms to map to the schema (e.g., "Job Title" or "Project Description" for "project_name", "Place" or "Site" for "project_location"). If a field is missing, infer it based on context or return an empty string. For quantities, handle multiple materials (e.g., asphalt, concrete) and convert units if necessary (e.g., ft³ to yd³ or tons).

Respond with a JSON object containing these keys:
- project_name (string)
- project_type (string, e.g., 'road', 'sidewalk', 'general')
- project_location (string)
- completion_date (string, format 'YYYY-MM-DD')
- project_duration (string, in weeks)
- land_mile (string, lane miles)
- width (string, in feet)
- project_area (string, in square feet)
- material_type (string, primary material, e.g., 'asphalt', 'concrete')
- tonnage (string, total tonnage for asphalt or aggregate)
- project_scope (string)
- project_requirements (string)
- quantities (array of objects with 'material', 'quantity', 'unit')

Example:
Text: "RFP #4: Residential Driveway & Sidewalk Replacement
PROJECT TITLE: Fox Hollow Estates – Driveway & Sidewalk Rehabilitation
PROJECT LOCATION: Fox Hollow Estates Subdivision, Lot #15–#28, Boulder City, NV
SCHEDULE: All work completed by August 15, 2025
ESTIMATED QUANTITIES: HMA driveway: 3,200 ft² × 0.333 ft ≈ 40 yd³; Concrete sidewalk: 1,800 ft² × 0.333 ft ≈ 22 yd³
SCOPE OF WORK: Remove existing driveway pavement; install 4” HMA surface..."

Output:
{
  "project_name": "Fox Hollow Estates – Driveway & Sidewalk Rehabilitation",
  "project_type": "sidewalk",
  "project_location": "Fox Hollow Estates Subdivision, Lot #15–#28, Boulder City, NV",
  "completion_date": "2025-08-15",
  "project_duration": "6",
  "land_mile": "",
  "width": "",
  "project_area": "5000",
  "material_type": "asphalt",
  "tonnage": "80",
  "project_scope": "Remove existing driveway pavement; install 4” HMA surface...",
  "project_requirements": "",
  "quantities": [
    {"material": "asphalt", "quantity": 80, "unit": "tons"},
    {"material": "concrete", "quantity": 22, "unit": "yd³"}
  ]
}

Text:
\"\"\"%s\"\"\"

Return the JSON object. Ensure dates are in 'YYYY-MM-DD' format. For project_type, infer from keywords (e.g., 'driveway' or 'sidewalk' implies 'sidewalk', 'lane' implies 'road'). If quantities are in ft³, convert to yd³ (divide by 27) or tons (use 150 lbs/ft³ for asphalt/concrete, 2000 lbs/ton). Limit scope and requirements to 1000 characters each.
""" % text[:3500]

    try:
        # Call OpenAI API
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2,
            max_tokens=1500
        )
        content = response.choices[0].message.content.strip()
        json_match = re.search(r'```json\n(.*?)\n```', content, re.DOTALL)
        if json_match:
            content = json_match.group(1)
        data = json.loads(content)
        
        # Validate and clean data
        for key in ['project_name', 'project_location', 'project_scope', 'project_requirements']:
            if key in data and data[key]:
                data[key] = data[key][:255] if key in ['project_name', 'project_location'] else data[key][:1000]
        if 'completion_date' in data and data['completion_date']:
            try:
                data['completion_date'] = parse_date(data['completion_date']).strftime('%Y-%m-%d')
            except:
                data['completion_date'] = ''
        
        return data
    except Exception as e:
        return {}

# Handle RFP file upload
@app.route('/upload_rfp', methods=['POST'])
def upload_rfp():
    """Process uploaded RFP files (PDF or DOCX) and generate project estimates."""
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400

    try:
        filename = file.filename.lower()
        file_data = file.read()
        file_stream = BytesIO(file_data)
        # Extract text based on file type
        if filename.endswith('.pdf'):
            text = extract_text_from_pdf(file_stream)
        elif filename.endswith('.docx'):
            text = extract_text_from_docx(file_stream)
        else:
            return jsonify({'error': 'Unsupported file type'}), 400

        # Try OpenAI GPT extraction first, fall back to regex
        extracted_data = extract_fields_with_openai(text)
        if not extracted_data:
            extracted_data = extract_rfp_data(text)

        # Set default values for missing fields
        if not extracted_data.get('project_name'):
            extracted_data['project_name'] = f"Project from {file.filename}"
        if not extracted_data.get('project_type'):
            extracted_data['project_type'] = 'road'
        if not extracted_data.get('project_location'):
            extracted_data['project_location'] = 'Unknown Location'
        if not extracted_data.get('project_scope'):
            extracted_data['project_scope'] = 'Scope not extracted'

        # Calculate area if not provided
        if not extracted_data.get('project_area') and extracted_data.get('land_mile') and extracted_data.get('width'):
            try:
                land_mile = float(extracted_data['land_mile'])
                width_ft = float(extracted_data['width'])
                if land_mile > 0 and width_ft > 0:
                    extracted_data['project_area'] = str(round((land_mile * 5280) * width_ft))
            except (ValueError, TypeError):
                pass

        if not extracted_data.get('project_area'):
            # Sum areas from quantities if available
            total_area = 0
            if extracted_data.get('quantities'):
                for q in extracted_data['quantities']:
                    if q['unit'] in ['ft²', 'sq ft'] and q['quantity'] > 0:
                        total_area += float(q['quantity'])
                if total_area > 0:
                    extracted_data['project_area'] = str(total_area)

        if not extracted_data.get('project_area'):
            return jsonify({
                'error': 'Could not determine project area. Please provide area or land-mile+width in the document.',
                'extracted_data': extracted_data
            }), 400

        # Prepare data for processing
        data = {
            'project_name': extracted_data.get('project_name'),
            'project_type': extracted_data.get('project_type', 'road'),
            'project_location': extracted_data.get('project_location'),
            'project_duration': extracted_data.get('project_duration', ''),
            'completion_date': extracted_data.get('completion_date', ''),
            'land_mile': extracted_data.get('land_mile', ''),
            'width': extracted_data.get('width', ''),
            'project_area': extracted_data.get('project_area'),
            'material_type': extracted_data.get('material_type', 'asphalt'),
            'tonnage': extracted_data.get('tonnage', ''),
            'project_scope': extracted_data.get('project_scope'),
            'project_requirements': extracted_data.get('project_requirements', ''),
            'quantities': extracted_data.get('quantities', [])
        }

        return process_estimate(data)

    except Exception as e:
        return jsonify({
            'error': 'RFP processing failed',
            'details': str(e)
        }), 500


# Process project estimate
def process_estimate(data):
    """Generate project estimate based on input data, including labor, materials, and financials."""

    # Helper function to safely convert values to float
    def safe_float(value, default=0.0):
        if not value or not str(value).strip() or str(value).lower() == "undefined":
            return default
        try:
            cleaned = re.sub(r'[^\d.\-]', '', str(value))
            return float(cleaned)
        except ValueError:
            return default

    try:
        # Extract and validate input data
        project_name = data.get('project_name', 'Unnamed Project')
        project_type = data.get('project_type', 'road')
        location = data.get('project_location', 'Unknown Location')
        scope = data.get('project_scope', '')
        project_requirements = data.get('project_requirements', '')
        material_type = data.get('material_type', 'asphalt').lower()
        tonnage = safe_float(data.get('tonnage', 0))
        
        land_mile = safe_float(data.get('land_mile', 0))
        width_ft = safe_float(data.get('width', 0))
        area_sqft = safe_float(data.get('project_area', 0))
        
        # Calculate area if not provided
        if area_sqft <= 0:
            if land_mile > 0 and width_ft > 0:
                area_sqft = (land_mile * 5280) * width_ft
            else:
                return jsonify({
                    'error': 'Valid area required: Provide either area or land-mile+width',
                    'details': f"land_mile: {land_mile}, width: {width_ft}, calculated_area: {area_sqft}"
                }), 400
        
        # Set default material type if not recognized
        if material_type not in ['asphalt', 'recycled asphalt', 'bituminous surface', 'concrete', 'sealcoat']:
            material_type = 'asphalt'
        
        # Determine project duration and completion date
        completion_date_str = data.get('completion_date', '')
        duration_weeks = safe_float(data.get('project_duration', 0))
        
        if completion_date_str:
            try:
                completion_date = datetime.strptime(completion_date_str, '%Y-%m-%d')
                today = datetime.now()
                duration_weeks = max((completion_date - today).days / 7, 1)
            except:
                completion_date = datetime.now() + timedelta(weeks=8)
                duration_weeks = 8
        else:
            if duration_weeks <= 0:
                duration_weeks = 8  # Default duration
            completion_date = datetime.now() + timedelta(weeks=duration_weeks)
        
        # Calculate estimates with detailed logging
        material_estimates = calculate_materials(area_sqft, material_type, tonnage)
        
        labor_estimates = calculate_labor(area_sqft, duration_weeks, project_type, material_type, width_ft)
        
        equipment_estimates = calculate_equipment(area_sqft, duration_weeks)
        
        financial_summary = calculate_financials(
            material_estimates, 
            labor_estimates, 
            equipment_estimates,
            area_sqft,
            duration_weeks,
            material_type
        )
        
        # Prepare project summary
        project_summary = {
            'project_name': project_name,
            'project_type': project_type.capitalize(),
            'location': location,
            'completion_date': completion_date.strftime('%Y-%m-%d'),
            'duration_weeks': duration_weeks,
            'area_sqft': round(area_sqft),
            'material_type': material_type.capitalize(),
            'land_mile': land_mile,
            'width': width_ft,
            'tonnage': tonnage 
        }
        
        success_probability = calculate_success_probability(project_type, area_sqft, duration_weeks)
        
        # Create new project record
        new_project = Project(
            name=project_name,
            type=project_type.capitalize(),
            location=location,
            submitted=datetime.now().date(),
            status='pending',
            cost=f"${financial_summary['total_cost']}",
            completion_date=completion_date.date(),
            land_mile=land_mile,
            width=width_ft,
            area=area_sqft,
            material=material_type.capitalize(),
            tonnage=tonnage if tonnage > 0 else material_estimates.get('asphalt_tons', material_estimates.get('concrete_yds', 0)),
            scope=scope,
            requirements=project_requirements,
            estimated_cost=f"${financial_summary['total_cost']}",
            profit_margin=financial_summary['profit_margin'],
            success_probability=success_probability,
            asphalt_tons=material_estimates.get('asphalt_tons', 0),
            concrete_yds=material_estimates.get('concrete_yds', 0),
            bituminous_tons=material_estimates.get('bituminous_tons', 0),
            sealcoat_sqft=material_estimates.get('sealcoat_sqft', 0),
            rebar_lbs=material_estimates.get('rebar_lbs', 0),
            aggregate_tons=material_estimates.get('aggregate_tons', 0),
            management_hours=labor_estimates.get('management_hours', 0),
            prep_hours=labor_estimates.get('prep_hours', 0),
            paving_hours=labor_estimates.get('paving_hours', 0),
            finishing_hours=labor_estimates.get('finishing_hours', 0)
        )
        
        # Save to database
        try:
            db.session.add(new_project)
            db.session.commit()
        except Exception as e:
            db.session.rollback()
            return jsonify({
                'error': 'Database operation failed',
                'details': str(e),
            }), 500

        # Prepare response
        response = {
            'project_summary': project_summary,
            'material_estimates': material_estimates,
            'labor_estimates': labor_estimates,
            'equipment_estimates': equipment_estimates,
            'financial_summary': financial_summary,
            'success_probability': success_probability,
            'project_id': new_project.id
        }
        
        return jsonify(response), 200
    
    except Exception as e:
        return jsonify({
            'error': 'Estimate processing failed',
            'details': str(e),
        }), 500
    

# Handle manual estimate calculation via JSON
@app.route('/calculate_estimate', methods=['POST'])
def calculate_estimate():
    data = request.json
    return process_estimate(data)

# Calculate material quantities
def calculate_materials(area_sqft, material_type, tonnage):
    try:
        material_type = material_type.lower()
        results = {}
        
        # Asphalt types
        if material_type in ['asphalt', 'recycled asphalt']:
            density = RECYCLED_DENSITY if 'recycled' in material_type else ASPHALT_DENSITY
            thickness = THICKNESS.get(material_type, THICKNESS['asphalt'])
            
            if tonnage > 0:
                asphalt_tons = tonnage
            else:
                volume_cf = area_sqft * thickness
                asphalt_tons = (volume_cf * density) / 2000
            
            results['asphalt_tons'] = round(asphalt_tons, 1)
            results['aggregate_tons'] = round(asphalt_tons * 1.2, 1)
        
        # Bituminous surface
        elif material_type == 'bituminous surface':
            density = MATERIAL_CONSTANTS['bituminous surface']['density']
            thickness = MATERIAL_CONSTANTS['bituminous surface']['thickness']
            
            if tonnage > 0:
                bituminous_tons = tonnage
            else:
                volume_cf = area_sqft * thickness
                bituminous_tons = (volume_cf * density) / 2000
            
            results['bituminous_tons'] = round(bituminous_tons, 1)
            results['aggregate_tons'] = round(bituminous_tons * 1.2, 1)
        
        # Concrete
        elif material_type == 'concrete':
            thickness = THICKNESS['concrete']
            volume_cf = area_sqft * thickness
            concrete_yds = volume_cf / 27
            results['concrete_yds'] = round(concrete_yds, 1)
            results['rebar_lbs'] = round(area_sqft * 1.2, 1)  # 1.2 lbs per sq ft
        
        # Sealcoat
        elif material_type == 'sealcoat':
            results['sealcoat_sqft'] = round(area_sqft)
        
        return results
    
    except Exception as e:
        # Just raise the exception without logging
        raise

# Calculate labor hours
def calculate_labor(area_sqft, duration_weeks, project_type, material_type, width_ft):
    """Calculate labor hours for a 7-person crew based on Virginia productivity rates."""
    try:
        # Set default width
        width_ft = width_ft or 0
        
        # Determine if project is a narrow path (width ≤ 3 ft)
        is_narrow = width_ft > 0 and width_ft <= 3
        
        # Set productivity rate based on project type and material
        if "road" in project_type.lower():
            if is_narrow and "concrete" in material_type.lower():
                sqft_per_crew_hour = 300  # Higher rate for narrow concrete paths
            else:
                sqft_per_crew_hour = 200  # Standard rate for asphalt/concrete roads
        elif "sidewalk" in project_type.lower():
            sqft_per_crew_hour = 150  # Rate for detailed sidewalk work
        else:
            sqft_per_crew_hour = 120  # General paving projects
        
        # Define crew size and weekly capacity
        CREW_SIZE = 7  # Standard crew size for small Virginia projects
        HOURS_PER_WORKER_PER_WEEK = 40  # Standard work week
        max_weekly_hours = CREW_SIZE * HOURS_PER_WORKER_PER_WEEK  # 280 hours/week
        
        # Calculate total crew hours
        total_crew_hours = area_sqft / sqft_per_crew_hour
        
        # Cap hours based on crew capacity and duration
        if duration_weeks > 0:
            max_total_hours = max_weekly_hours * duration_weeks
            total_crew_hours = min(total_crew_hours, max_total_hours)
        
        # Ensure minimum hours for small projects
        total_crew_hours = max(total_crew_hours, CREW_SIZE * 8)  # Minimum 56 hours
        
        # Set phase distribution percentages
        if is_narrow:
            management_pct = 0.10  # 10% for management
            prep_pct = 0.20       # 20% for minimal site prep
            paving_pct = 0.65     # 65% for main paving work
            finishing_pct = 0.05  # 5% for minimal finishing
        else:
            management_pct = 0.10  # Standard VDOT distribution
            prep_pct = 0.30
            paving_pct = 0.50
            finishing_pct = 0.10
        
        # Distribute hours across phases
        management_hours = total_crew_hours * management_pct
        prep_hours = total_crew_hours * prep_pct
        paving_hours = total_crew_hours * paving_pct
        finishing_hours = total_crew_hours * finishing_pct
        
        # Return rounded hours
        return {
            'management_hours': round(management_hours),
            'prep_hours': round(prep_hours),
            'paving_hours': round(paving_hours),
            'finishing_hours': round(finishing_hours),
            'total_hours': round(total_crew_hours)
        }
    
    except Exception as e:
        raise

# Calculate equipment needs
def calculate_equipment(area_sqft, duration_weeks):
    """Calculate equipment quantities and costs based on Virginia rental rates."""
    try:
        pavers = max(1, math.ceil(area_sqft / 120000))
        rollers = max(1, math.ceil(area_sqft / 60000))
        excavators = max(1, math.ceil(area_sqft / 150000))
        trucks = max(2, math.ceil(area_sqft / 50000))
        
        paver_cost = pavers * 2500 * duration_weeks
        roller_cost = rollers * 1000 * duration_weeks
        excavator_cost = excavators * 2000 * duration_weeks
        truck_cost = trucks * 900 * duration_weeks

        return {
            'pavers': pavers,
            'rollers': rollers,
            'excavators': excavators,
            'trucks': trucks,
            'paver_cost': round(paver_cost),
            'roller_cost': round(roller_cost),
            'excavator_cost': round(excavator_cost),
            'truck_cost': round(truck_cost)
        }
    
    except Exception as e:
        raise


# Calculate financial summary
def calculate_financials(materials, labor, equipment, area_sqft, duration_weeks, material_type):
    """Calculate project costs, including materials, labor, equipment, overhead, and profit."""
    try:
        material_type = material_type.lower()
        material_costs = 0

        # Calculate material costs
        if material_type in ['asphalt', 'recycled asphalt', 'bituminous surface']:
            # Determine asphalt unit cost
            if material_type == 'recycled asphalt':
                unit_cost = MATERIAL_UNIT_COSTS['recycled asphalt']
            elif material_type == 'bituminous surface':
                unit_cost = MATERIAL_UNIT_COSTS['bituminous surface']
            else:
                unit_cost = MATERIAL_UNIT_COSTS['asphalt']
            
            # Calculate asphalt costs
            asphalt_tons = materials.get('asphalt_tons', 0)
            material_costs += asphalt_tons * unit_cost * MATERIAL_MARKUP
            
            # Calculate aggregate costs
            aggregate_tons = materials.get('aggregate_tons', 0)
            material_costs += aggregate_tons * MATERIAL_UNIT_COSTS['aggregate base'] * MATERIAL_MARKUP
        
        elif material_type == 'concrete':
            # Calculate concrete costs
            concrete_yds = materials.get('concrete_yds', 0)
            material_costs += concrete_yds * MATERIAL_UNIT_COSTS['concrete'] * MATERIAL_MARKUP
            
            # Calculate rebar costs
            rebar_lbs = materials.get('rebar_lbs', 0)
            material_costs += rebar_lbs * MATERIAL_UNIT_COSTS['rebar'] * MATERIAL_MARKUP
        
        elif material_type == 'sealcoat':
            # Calculate sealcoat costs
            sealcoat_sqft = materials.get('sealcoat_sqft', area_sqft)
            material_costs += sealcoat_sqft * MATERIAL_UNIT_COSTS['sealcoat'] * MATERIAL_MARKUP

        # Calculate labor costs
        labor_hours = labor.get('total_hours', 0)
        if labor_hours <= 0:
            # Fallback calculation if labor hours not provided
            labor_hours = area_sqft / 100 * 10  # 10 hours per 100 sq ft
        labor_costs = labor_hours * LABOR_RATE
        
        # Calculate equipment costs with markup
        equipment_costs = (
            equipment.get('paver_cost', 0) + 
            equipment.get('roller_cost', 0) + 
            equipment.get('excavator_cost', 0) + 
            equipment.get('truck_cost', 0)
        ) * EQUIPMENT_RATE_MULTIPLIER
        
        # Calculate subtotal, overhead, and profit
        subtotal = material_costs + labor_costs + equipment_costs
        overhead = subtotal * OVERHEAD_RATE
        profit = subtotal * PROFIT_MARGIN
        total_cost = subtotal + overhead + profit
        
        # Prepare cost breakdown
        cost_breakdown = {
            'materials': round(material_costs),
            'labor': round(labor_costs),
            'equipment': round(equipment_costs),
            'overhead': round(overhead),
            'profit': round(profit)
        }
        
        cost_per_sqft = total_cost / area_sqft if area_sqft > 0 else 0
        
        return {
            'total_cost': round(total_cost),
            'cost_per_sqft': round(cost_per_sqft, 2),
            'profit_margin': f"{PROFIT_MARGIN * 100}%",
            'cost_breakdown': cost_breakdown
        }
    
    except Exception as e:
        raise


# Calculate bid success probability
def calculate_success_probability(project_type, area_sqft, duration_weeks):
    """Estimate probability of bid success based on project factors."""
    base_prob = 75
    
    if project_type.lower() == 'road':
        base_prob += 5
    elif project_type.lower() == 'renovation':
        base_prob -= 5
    
    if area_sqft > 150000:
        base_prob -= 8
    elif area_sqft < 15000:
        base_prob += 5
    
    if duration_weeks > 24:
        base_prob -= 7
    elif duration_weeks < 6:
        base_prob += 5
    
    probability = max(60, min(95, base_prob))
    
    return f"{probability}%"


# Download project report as PDF
@app.route('/download_report/<int:project_id>', methods=['GET'])
def download_report(project_id):
    project = db.session.get(Project, project_id)
    if not project:
        return jsonify({'error': 'Project not found'}), 404
    
    pdf = generate_pdf_report(project)
    
    response = make_response(pdf)
    response.headers['Content-Type'] = 'application/pdf'
    response.headers['Content-Disposition'] = f'attachment; filename=project_{project_id}_report.pdf'
    return response


# Download project report as CSV
@app.route('/download_report_csv/<int:project_id>', methods=['GET'])
def download_report_csv(project_id):
    project = db.session.get(Project, project_id)
    if not project:
        return jsonify({'error': 'Project not found'}), 404
    
    si = StringIO()
    cw = csv.writer(si)
    
    # Write CSV content
    cw.writerow(['Project Report', f'Project ID: {project_id}'])
    cw.writerow([])
    cw.writerow(['Field', 'Value'])
    cw.writerow(['Project Name', project.name])
    cw.writerow(['Project Type', project.type])
    cw.writerow(['Location', project.location])
    cw.writerow(['Submitted Date', project.submitted.strftime('%Y-%m-%d')])
    cw.writerow(['Status', project.status])
    cw.writerow(['Estimated Cost', project.cost])
    cw.writerow(['Completion Date', project.completion_date.strftime('%Y-%m-%d') if project.completion_date else ''])
    cw.writerow(['Land-Mile', project.land_mile])
    cw.writerow(['Width (ft)', project.width])
    cw.writerow(['Area (sq ft)', project.area])
    
    
    material_type = project.material.lower()
    cw.writerow([])
    cw.writerow(['Material Estimates', 'Quantity'])
    
    if material_type == 'bituminous surface':
        if hasattr(project, 'bituminous_tons') and project.bituminous_tons:
            cw.writerow(['Bituminous Surface', f'{project.bituminous_tons} tons'])
        if project.aggregate_tons:
            cw.writerow(['Aggregate', f'{project.aggregate_tons} tons'])
    elif 'asphalt' in material_type or 'bituminous' in material_type:
        if project.asphalt_tons:
            cw.writerow(['Asphalt', f'{project.asphalt_tons} tons'])
        if project.aggregate_tons:
            cw.writerow(['Aggregate', f'{project.aggregate_tons} tons'])
    elif 'concrete' in material_type:
        if project.concrete_yds:
            cw.writerow(['Concrete', f'{project.concrete_yds} cubic yards'])
        if project.rebar_lbs:
            cw.writerow(['Rebar', f'{project.rebar_lbs} lbs'])
    elif material_type == 'sealcoat':
        if project.sealcoat_sqft:
            cw.writerow(['Sealcoat', f'{project.sealcoat_sqft} sq ft'])

    cw.writerow(['Management Hours', project.management_hours])
    cw.writerow(['Preparation Hours', project.prep_hours])
    cw.writerow(['Paving Hours', project.paving_hours])
    cw.writerow(['Finishing Hours', project.finishing_hours])
    cw.writerow(['Profit Margin', project.profit_margin])
    cw.writerow(['Success Probability', project.success_probability])
    cw.writerow(['Scope', project.scope])
    cw.writerow(['Requirements', project.requirements or ''])
    
    response = make_response(si.getvalue())
    response.headers['Content-Type'] = 'text/csv'
    response.headers['Content-Disposition'] = f'attachment; filename=project_{project_id}_report.csv'
    return response


# Generate PDF report
def generate_pdf_report(project):
    """Generate a styled PDF report for a project."""
    current_date = datetime.now().strftime('%B %d, %Y')
    logo_path = os.path.join(app.root_path, 'static', 'images', 'logo.png')
    logo_data = ""
    if os.path.exists(logo_path):
        with open(logo_path, "rb") as logo_file:
            logo_data = base64.b64encode(logo_file.read()).decode('utf-8')
    
    material_rows = ""
    material_type = project.material.lower()
    
    if material_type == 'bituminous surface':
        if hasattr(project, 'bituminous_tons') and project.bituminous_tons:
            material_rows += f"<tr><td>Bituminous Surface</td><td>{project.bituminous_tons} tons</td></tr>"
        if project.aggregate_tons:
            material_rows += f"<tr><td>Aggregate</td><td>{project.aggregate_tons} tons</td></tr>"
    elif 'asphalt' in material_type or 'bituminous' in material_type:
        if project.asphalt_tons:
            material_rows += f"<tr><td>Asphalt</td><td>{project.asphalt_tons} tons</td></tr>"
        if project.aggregate_tons:
            material_rows += f"<tr><td>Aggregate</td><td>{project.aggregate_tons} tons</td></tr>"
    elif 'concrete' in material_type:
        if project.concrete_yds:
            material_rows += f"<tr><td>Concrete</td><td>{project.concrete_yds} cubic yards</td></tr>"
        if project.rebar_lbs:
            material_rows += f"<tr><td>Rebar</td><td>{project.rebar_lbs} lbs</td></tr>"
    elif material_type == 'sealcoat':
        if project.sealcoat_sqft:
            material_rows += f"<tr><td>Sealcoat</td><td>{project.sealcoat_sqft} sq ft</td></tr>"


    # Define HTML content for PDF
    html_content = f"""
    <html>
    <head>
        <title>Project Report - {project.id}</title>
        <style>
            @page {{ size: A4; margin: 1.5cm; }}
            body {{ font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; color: #333; line-height: 1.6; }}
            .header {{ border-bottom: 2px solid #3498db; padding-bottom: 15px; margin-bottom: 25px; }}
            h1 {{ color: #2c3e50; margin-bottom: 5px; }}
            h2 {{ color: #3498db; border-bottom: 1px solid #eee; padding-bottom: 8px; margin-top: 25px; }}
            .subtitle {{ color: #7f8c8d; font-size: 1.1rem; }}
            .project-info {{
                background-color: #f8f9fa;
                border-radius: 8px;
                padding: 25px;
                margin-bottom: 35px;
                box-shadow: 0 2px 4px rgba(0, 0, 0, 0.05);
            }}
            .grid-container {{ display: grid; grid-template-columns: 1fr 1fr; gap: 30px; }}
            .grid-container p {{
                margin: 10px 0;
            }}
            .section {{ margin-bottom: 25px; }}
            table {{ width: 100%; border-collapse: collapse; margin: 15px 0; }}
            th {{ background-color: #3498db; color: white; text-align: left; padding: 12px; }}
            td {{ padding: 10px; border-bottom: 1px solid #eee; }}
            tr:nth-child(even) {{ background-color: #f9f9f9; }}
            .footer {{ margin-top: 40px; text-align: center; color: #7f8c8d; font-size: 0.9rem; }}
            .status-badge {{
                display: inline-block;
                padding: 3px 10px;
                border-radius: 12px;
                font-size: 0.85rem;
                font-weight: bold;
                margin-left: 10px;
            }}
            .status-pending {{ background-color: #f39c12; color: white; }}
            .status-accepted {{ background-color: #27ae60; color: white; }}
            .status-rejected {{ background-color: #e74c3c; color: white; }}
            .header-logo-container {{ 
                display: flex;
                justify-content: center;
                margin-bottom: 10px;
            }}
            .logo-img {{ 
                max-width: 100%; 
                height: auto; 
            }}
            .logo {{ font-weight: bold; font-size: 1.8rem; margin-bottom: 5px; }}
            .logo span {{ color: #f39c12; }}
        </style>
    </head>
    <body>
        <div class="header-logo-container">
            {"<img src='data:image/png;base64," + logo_data + "' class='logo-img'/>" if logo_data else ""}
        </div>
        <div class="header">
            <div class="logo">Bid<span>Master</span></div>
            <h1>Project Report: {project.name}</h1>
            <p class="subtitle">Generated on {current_date} | Project ID: {project.id}</p>
        </div>
        <div class="project-info">
            <div class="grid-container">
                <div>
                    <p><strong>Project Type:</strong> {project.type}</p>
                    <p><strong>Location:</strong> {project.location}</p>
                    <p><strong>Land-mile:</strong> {project.land_mile} mile</p>
                    <p><strong>Width:</strong> {project.width} ft</p>
                    <p><strong>Area:</strong> {project.area} sq ft</p>
                    <p><strong>Material:</strong> {project.material}</p>
                </div>
                <div>
                    <p><strong>Estimated Cost:</strong> {project.cost}</p>
                    <p><strong>Submitted on:</strong> {project.submitted.strftime('%Y-%m-%d')}</p>
                    <p><strong>Completion Date:</strong> {project.completion_date.strftime('%Y-%m-%d') if project.completion_date else 'N/A'}</p>
                    <p><strong>Tonnage:</strong> {project.tonnage} tons</p>
                    <p><strong>Status:</strong> 
                        <span class="status-badge status-{project.status}">{project.status.capitalize()}</span>
                    </p>
                </div>
            </div>
        </div>
        <div class="section">
            <h2>Resource Estimates</h2>
            <table>
                <tr>
                    <th>Material</th>
                    <th>Quantity</th>
                </tr>
                {material_rows}
            </table>
        </div>
        <div class="section">
            <h2>Labor Estimates</h2>
            <table>
                <tr>
                    <th>Task</th>
                    <th>Hours</th>
                </tr>
                <tr>
                    <td>Management</td>
                    <td>{project.management_hours}</td>
                </tr>
                <tr>
                    <td>Preparation</td>
                    <td>{project.prep_hours}</td>
                </tr>
                <tr>
                    <td>Paving</td>
                    <td>{project.paving_hours}</td>
                </tr>
                <tr>
                    <td>Finishing</td>
                    <td>{project.finishing_hours}</td>
                </tr>
            </table>
        </div>
        <div class="section">
            <h2>Financial Summary</h2>
            <table>
                <tr>
                    <td>Estimated Cost:</td>
                    <td>{project.cost}</td>
                    <td>Profit Margin:</td>
                    <td>{project.profit_margin}</td>
                </tr>
                <tr>
                    <td>Success Probability:</td>
                    <td>{project.success_probability}</td>
                </tr>
            </table>
        </div>
        <div class="section">
            <h2>Project Scope</h2>
            <p>{project.scope}</p>
        </div>
        <div class="section">
            <h2>Requirements</h2>
            <p>{project.requirements or 'No special requirements specified'}</p>
        </div>
        <div class="footer">
            <p>Generated by Paveiq BidMaster System</p>
            <p>© {datetime.now().year} Paveiq. All rights reserved.</p>
        </div>
    </body>
    </html>
    """
    
    font_config = FontConfiguration()
    return HTML(string=html_content).write_pdf(font_config=font_config)


# Main: Run the Flask app
if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=False)
